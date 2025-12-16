"""
Routes cho tính năng tạo CV thông minh với AI
"""

import json
import logging
import os
from flask import Blueprint, request, jsonify, render_template
from werkzeug.utils import secure_filename
from app.extensions import db
from app.models import User, JobCategory
from app.models.cvdata import CVData
from flask_jwt_extended import jwt_required, get_jwt_identity
from app.utils.ai_enhancer import enhance_full_cv_with_ai, enhance_summary_with_ai, enhance_experience_with_ai, enhance_skills_with_ai
from app.utils.classifier import classify_cv_by_keywords, get_category_id_by_name
from datetime import datetime

logger = logging.getLogger(__name__)

cv_builder_bp = Blueprint('cv_builder', __name__)

# Thư mục lưu ảnh đại diện
AVATAR_UPLOAD_FOLDER = os.path.join('Flask_CVProject', 'app', 'static', 'uploads', 'avatars')
ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}

def allowed_image_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_IMAGE_EXTENSIONS

def save_avatar(file, user_id):
    """Lưu ảnh đại diện và trả về đường dẫn"""
    if file and allowed_image_file(file.filename):
        os.makedirs(AVATAR_UPLOAD_FOLDER, exist_ok=True)
        file_ext = file.filename.rsplit('.', 1)[1].lower()
        filename = f'avatar_{user_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.{file_ext}'
        filename = secure_filename(filename)
        file_path = os.path.join(AVATAR_UPLOAD_FOLDER, filename)
        file.save(file_path)
        return f'/static/uploads/avatars/{filename}'
    return None

def get_user_id_from_jwt():
    """Helper function to get user_id from JWT identity and convert to int"""
    user_id = get_jwt_identity()
    if isinstance(user_id, str) and user_id.isdigit():
        return int(user_id)
    return user_id

# Route lưu CV đã bị xóa - không còn cần thiết

@cv_builder_bp.route('/api/cv-builder/export-temp', methods=['POST'])
@jwt_required()
def export_temp_cv():
    """Export CV trực tiếp từ form mà không cần lưu DB"""
    try:
        data = request.get_json()
        template = data.get('template', 'classic')
        format_type = data.get('format', 'pdf')  # pdf hoặc docx

        # Process data same as preview
        experiences = data.get('experiences', [])
        processed_experiences = []
        for exp in experiences:
            processed_exp = {
                'position': exp.get('position', ''),
                'company': exp.get('company', ''),
                'start_date': exp.get('start_date', ''),
                'end_date': exp.get('end_date', ''),
                'is_current': exp.get('is_current', False),
                'description': exp.get('description', '')
            }
            if processed_exp['end_date'] == '' and processed_exp['is_current']:
                processed_exp['end_date'] = None
            processed_experiences.append(processed_exp)

        education = data.get('education', [])
        processed_education = []
        for edu in education:
            processed_edu = {
                'school': edu.get('school', ''),
                'major': edu.get('major', ''),
                'degree': edu.get('degree', ''),
                'year': edu.get('year', '')
            }
            processed_education.append(processed_edu)

        skills = data.get('skills', [])
        if isinstance(skills, str):
            skills = [s.strip() for s in skills.split(',') if s.strip()]
        elif not isinstance(skills, list):
            skills = []

        certifications = data.get('certifications', [])
        processed_certifications = []
        for cert in certifications:
            processed_cert = {
                'name': cert.get('name', ''),
                'organization': cert.get('organization', ''),
                'date': cert.get('date', '')
            }
            processed_certifications.append(processed_cert)

        projects = data.get('projects', [])
        processed_projects = []
        for proj in projects:
            processed_proj = {
                'name': proj.get('name', ''),
                'url': proj.get('url', ''),
                'description': proj.get('description', '')
            }
            processed_projects.append(processed_proj)

        languages = data.get('languages', [])
        processed_languages = []
        for lang in languages:
            processed_lang = {
                'name': lang.get('name', ''),
                'level': lang.get('level', '')
            }
            processed_languages.append(processed_lang)

        # Xử lý upload ảnh nếu có
        avatar_url = data.get('avatar_url', '')
        if 'avatar' in request.files:
            user_id = get_user_id_from_jwt()
            file = request.files['avatar']
            if file and file.filename != '':
                avatar_path = save_avatar(file, user_id)
                if avatar_path:
                    avatar_url = avatar_path
        
        # Import hàm xử lý UTF-8
        try:
            from app.routes.recruiter_routes import safe_decode_text, fix_common_vietnamese_errors
        except ImportError:
            # Fallback nếu không import được
            def safe_decode_text(text):
                if isinstance(text, bytes):
                    return text.decode('utf-8', errors='replace')
                return str(text) if text else ''
            def fix_common_vietnamese_errors(text):
                return text
        
        # Đảm bảo tất cả dữ liệu text được decode và fix UTF-8 đúng cách
        def ensure_utf8(text):
            if not text:
                return text
            text = safe_decode_text(text)
            text = fix_common_vietnamese_errors(text)
            return text
        
        # Xử lý tất cả text fields
        full_name = ensure_utf8(data.get('full_name', ''))
        email = ensure_utf8(data.get('email', ''))
        phone = ensure_utf8(data.get('phone', ''))
        address = ensure_utf8(data.get('address', ''))
        linkedin = ensure_utf8(data.get('linkedin', ''))
        website = ensure_utf8(data.get('website', ''))
        summary = ensure_utf8(data.get('summary', ''))
        
        # Xử lý experiences
        for exp in processed_experiences:
            exp['position'] = ensure_utf8(exp.get('position', ''))
            exp['company'] = ensure_utf8(exp.get('company', ''))
            exp['description'] = ensure_utf8(exp.get('description', ''))
        
        # Xử lý education
        for edu in processed_education:
            edu['school'] = ensure_utf8(edu.get('school', ''))
            edu['major'] = ensure_utf8(edu.get('major', ''))
            edu['degree'] = ensure_utf8(edu.get('degree', ''))
        
        # Xử lý skills - đảm bảo filter empty strings và ensure UTF-8
        if isinstance(skills, list):
            skills = [ensure_utf8(str(skill).strip()) for skill in skills if skill and str(skill).strip()]
        else:
            skills = []
        
        # Xử lý certifications
        for cert in processed_certifications:
            cert['name'] = ensure_utf8(cert.get('name', ''))
            cert['organization'] = ensure_utf8(cert.get('organization', ''))
        
        # Xử lý projects
        for proj in processed_projects:
            proj['name'] = ensure_utf8(proj.get('name', ''))
            proj['description'] = ensure_utf8(proj.get('description', ''))
        
        # Xử lý languages
        for lang in processed_languages:
            lang['name'] = ensure_utf8(lang.get('name', ''))
            lang['level'] = ensure_utf8(lang.get('level', ''))
        
        # Đảm bảo không có giá trị None trong context để tránh lỗi khi tạo PDF
        context = {
            'full_name': full_name or '',
            'email': email or '',
            'phone': phone or '',
            'address': address or '',
            'linkedin': linkedin or '',
            'website': website or '',
            'summary': summary or '',
            'avatar_url': avatar_url or '',
            'experiences': processed_experiences or [],
            'education': processed_education or [],
            'skills': skills or [],
            'certifications': processed_certifications or [],
            'projects': processed_projects or [],
            'languages': processed_languages or []
        }

        template_file = f"cv_formats/{template}.html"
        cv_html = render_template(template_file, **context)

        if format_type == 'pdf':
           
            import io
            import base64
            import re
            import unicodedata
            import tempfile
            import os
            
            html_content = cv_html
            
            html_content = unicodedata.normalize('NFC', html_content)
            
            if not html_content.strip().startswith('<!DOCTYPE'):
                html_content = '<!DOCTYPE html>\n' + html_content
            if '<meta charset="UTF-8">' not in html_content:
                html_content = html_content.replace('<head>', '<head>\n    <meta charset="UTF-8">', 1)
            if '<meta http-equiv="Content-Type"' not in html_content:
                html_content = html_content.replace(
                    '<meta charset="UTF-8">',
                    '<meta charset="UTF-8">\n    <meta http-equiv="Content-Type" content="text/html; charset=UTF-8">',
                    1
                )
            
            def replace_image_src(match):
                src = match.group(1)
                if src.startswith('/static/'):
                    img_path = os.path.join(os.getcwd(), 'Flask_CVProject', 'app', src.lstrip('/'))
                    if os.path.exists(img_path):
                        try:
                            with open(img_path, 'rb') as img_file:
                                img_data = img_file.read()
                                img_base64 = base64.b64encode(img_data).decode('utf-8')
                                ext = os.path.splitext(img_path)[1].lower()
                                mime_type = {
                                    '.jpg': 'image/jpeg',
                                    '.jpeg': 'image/jpeg',
                                    '.png': 'image/png',
                                    '.webp': 'image/webp'
                                }.get(ext, 'image/jpeg')
                                return f'src="data:{mime_type};base64,{img_base64}"'
                        except Exception as e:
                            logger.warning(f"Could not encode image {img_path}: {e}")
                return match.group(0)
            
            html_content = re.sub(r'src="([^"]+)"', replace_image_src, html_content)
            
            try:
                from playwright.sync_api import sync_playwright
                import io
                
                pdf_stream = io.BytesIO()
                
                if isinstance(html_content, bytes):
                    html_content = html_content.decode('utf-8')
                
                with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as tmp_file:
                    tmp_file.write(html_content)
                    tmp_html_path = tmp_file.name
                
                try:
                    # Dùng Playwright để render HTML thành PDF (giống browser thật)
                    with sync_playwright() as p:
                        browser = p.chromium.launch(headless=True)
                        page = browser.new_page()
                        
                        # Load HTML từ file
                        page.goto(f'file:///{tmp_html_path.replace(os.sep, "/")}')
                        
                        # Chờ page load xong
                        page.wait_for_load_state('networkidle')
                        
                        # Tạo PDF với format A4
                        pdf_bytes = page.pdf(
                            format='A4',
                            margin={
                                'top': '20mm',
                                'right': '20mm',
                                'bottom': '20mm',
                                'left': '20mm'
                            },
                            print_background=True,  # Đảm bảo background colors được in
                            prefer_css_page_size=True
                        )
                        
                        browser.close()
                        
                        pdf_stream.write(pdf_bytes)
                        pdf_stream.seek(0)
                        
                        return (
                            pdf_stream.read(),
                            200,
                            {
                                "Content-Type": "application/pdf",
                                "Content-Disposition": f"attachment; filename={data.get('full_name','CV')}.pdf",
                            },
                        )
                finally:
                    # Xóa file tạm
                    try:
                        os.unlink(tmp_html_path)
                    except:
                        pass
                        
            except ImportError:
                logger.warning("Playwright chưa được cài đặt, thử weasyprint")
                # Fallback về weasyprint
                try:
                    from weasyprint import HTML, CSS
                    import io
                    
                    pdf_stream = io.BytesIO()
                    
                    # Đảm bảo HTML là string Unicode
                    if isinstance(html_content, bytes):
                        html_content = html_content.decode('utf-8')
                    
                    # Tạo base_url cho weasyprint để load resources
                    base_url_path = os.path.abspath(os.path.join('Flask_CVProject', 'app'))
                    if os.name == 'nt':
                        base_url = f'file:///{base_url_path.replace(os.sep, "/")}'
                    else:
                        base_url = f'file://{base_url_path}'
                    
                    # Tạo PDF với weasyprint - giữ nguyên CSS từ template
                    html_doc = HTML(string=html_content.encode('utf-8'), base_url=base_url)
                    html_doc.write_pdf(
                        pdf_stream,
                        stylesheets=[CSS(string='''
                            @page {
                                size: A4;
                                margin: 20mm;
                            }
                            /* Đảm bảo font hỗ trợ tiếng Việt */
                            body {
                                font-family: 'Times New Roman', 'Arial', 'Tahoma', sans-serif;
                            }
                            /* Đảm bảo flexbox hoạt động trong PDF */
                            .cv-header {
                                display: flex !important;
                                align-items: flex-start !important;
                                gap: 30px !important;
                            }
                            .cv-header .avatar-container {
                                flex-shrink: 0 !important;
                                width: 150px !important;
                                height: 150px !important;
                            }
                            .cv-header .header-content {
                                flex-grow: 1 !important;
                            }
                            .contact-info {
                                display: flex !important;
                                flex-wrap: wrap !important;
                                gap: 15px !important;
                            }
                        ''')]
                    )
                    
                    pdf_stream.seek(0)
                    return (
                        pdf_stream.read(),
                        200,
                        {
                            "Content-Type": "application/pdf",
                            "Content-Disposition": f"attachment; filename={data.get('full_name','CV')}.pdf",
                        },
                    )
                except (ImportError, OSError, Exception) as weasy_error:
                    error_msg = str(weasy_error)
                    if 'libgobject' in error_msg or 'GTK' in error_msg or 'OSError' in str(type(weasy_error)):
                        logger.warning(f"weasyprint không khả dụng (thiếu GTK+): {weasy_error}, thử xhtml2pdf với cải thiện")
                    else:
                        logger.warning(f"weasyprint không khả dụng: {weasy_error}, thử xhtml2pdf với cải thiện")
                    
                    # Fallback về xhtml2pdf với cải thiện UTF-8 và CSS
                try:
                    from xhtml2pdf import pisa
                    import io
                    
                    pdf_stream = io.BytesIO()
                    
                    # Đảm bảo HTML là string Unicode
                    if isinstance(html_content, bytes):
                        html_content = html_content.decode('utf-8')
                    
                    # Thêm CSS inline để xhtml2pdf hiểu được layout (vì xhtml2pdf không hỗ trợ flexbox tốt)
                    # Chuyển đổi flexbox thành table layout cho xhtml2pdf
                    # Loại bỏ các CSS properties không được hỗ trợ (như gap, object-fit)
                    css_fix = '''
                    <style type="text/css">
                        /* Fix cho xhtml2pdf - chuyển flexbox thành table để layout giống preview */
                        .cv-header {
                            display: table !important;
                            width: 100% !important;
                            border-collapse: separate !important;
                            border-spacing: 30px 0 !important;
                            margin-bottom: 30px !important;
                            padding-bottom: 20px !important;
                            border-bottom: 3px solid #2c3e50 !important;
                        }
                        .cv-header .avatar-container {
                            display: table-cell !important;
                            width: 150px !important;
                            vertical-align: top !important;
                        }
                        .cv-header .avatar-container img {
                            width: 150px !important;
                            height: 150px !important;
                            max-width: 150px !important;
                            max-height: 150px !important;
                        }
                        .cv-header .header-content {
                            display: table-cell !important;
                            vertical-align: top !important;
                            width: auto !important;
                        }
                        .cv-header h1 {
                            font-size: 32px !important;
                            font-weight: bold !important;
                            color: #2c3e50 !important;
                            margin-bottom: 10px !important;
                            text-transform: uppercase !important;
                            letter-spacing: 2px !important;
                        }
                        .contact-info {
                            display: block !important;
                            margin-top: 10px !important;
                        }
                        .contact-item {
                            display: inline-block !important;
                            margin-right: 15px !important;
                            margin-bottom: 5px !important;
                        }
                        .section-title {
                            font-size: 18px !important;
                            font-weight: bold !important;
                            color: #2c3e50 !important;
                            border-bottom: 2px solid #3498db !important;
                            padding-bottom: 5px !important;
                            margin-bottom: 15px !important;
                            text-transform: uppercase !important;
                            letter-spacing: 1px !important;
                        }
                        /* Đảm bảo font hỗ trợ tiếng Việt */
                        body, p, div, span, h1, h2, h3, h4, h5, h6, li, td, th {
                            font-family: 'Times New Roman', 'Arial', 'Tahoma', sans-serif !important;
                        }
                    </style>
                    '''
                    # Chỉ thêm CSS nếu có thẻ </head>
                    if '</head>' in html_content:
                        html_content = html_content.replace('</head>', css_fix + '</head>')
                    else:
                        # Nếu không có </head>, thêm vào sau <head>
                        if '<head>' in html_content:
                            html_content = html_content.replace('<head>', '<head>' + css_fix)
                        else:
                            # Nếu không có <head>, thêm vào đầu body
                            if '<body>' in html_content:
                                html_content = html_content.replace('<body>', css_fix + '<body>')
                            else:
                                html_content = css_fix + html_content
                    
                    # Đảm bảo HTML được encode đúng UTF-8
                    try:
                        html_content.encode('utf-8')
                    except UnicodeEncodeError:
                        html_content = html_content.encode('utf-8', errors='replace').decode('utf-8')
                    
                    # Loại bỏ các ký tự không hợp lệ có thể gây lỗi
                    # Xử lý các giá trị None trong HTML
                    html_content = html_content.replace('None', '')
                    html_content = html_content.replace('null', '')
                    
                    # Tạo PDF với xhtml2pdf - dùng string Unicode
                    try:
                        result = pisa.CreatePDF(
                            html_content,
                            dest=pdf_stream,
                            encoding='utf-8',
                            link_callback=None
                        )
                        
                        if result.err:
                            logger.error(f"xhtml2pdf error: {result.err}")
                            raise Exception(f"Lỗi khi tạo PDF: {result.err}")
                    except TypeError as type_err:
                        # Xử lý lỗi so sánh None với int
                        logger.error(f"TypeError trong xhtml2pdf: {type_err}")
                        # Thử lại với HTML đã được làm sạch hơn
                        import html as html_escape
                        # Escape các ký tự đặc biệt có thể gây lỗi
                        html_content_clean = html_content
                        # Loại bỏ các thuộc tính CSS có thể gây lỗi
                        html_content_clean = re.sub(r'gap:\s*[^;]+;', '', html_content_clean)
                        html_content_clean = re.sub(r'object-fit:\s*[^;]+;', '', html_content_clean)
                        
                        result = pisa.CreatePDF(
                            html_content_clean,
                            dest=pdf_stream,
                            encoding='utf-8',
                            link_callback=None
                        )
                        
                        if result.err:
                            logger.error(f"xhtml2pdf error sau khi clean: {result.err}")
                            raise Exception(f"Lỗi khi tạo PDF: {result.err}")
                    
                    pdf_stream.seek(0)
                    return (
                        pdf_stream.read(),
                        200,
                        {
                            "Content-Type": "application/pdf",
                            "Content-Disposition": f"attachment; filename={data.get('full_name','CV')}.pdf",
                        },
                    )
                except Exception as xhtml_error:
                    logger.error(f"xhtml2pdf cũng thất bại: {xhtml_error}", exc_info=True)
                    return jsonify({
                        'success': False,
                        'message': f'Lỗi khi tạo PDF: {str(xhtml_error)}. Vui lòng cài đặt weasyprint hoặc xhtml2pdf.'
                    }), 500

        elif format_type == 'docx':
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
            
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(11)
            
            # Header với ảnh nếu có
            if avatar_url and avatar_url.startswith('/static/'):
                avatar_path = os.path.join('Flask_CVProject', 'app', avatar_url.lstrip('/'))
                if os.path.exists(avatar_path):
                    try:
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run()
                        # Kích thước ảnh: 1.5 inch = ~150px (giống với preview)
                        run.add_picture(avatar_path, width=Inches(1.5), height=Inches(1.5))
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.space_after = Pt(12)
                    except Exception as e:
                        logger.warning(f"Could not add avatar to DOCX: {e}")
            
            # Tên
            heading = doc.add_heading(data.get('full_name', 'CV'), 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Thông tin liên hệ
            contact_info = []
            if data.get('email'): contact_info.append(f"Email: {data.get('email')}")
            if data.get('phone'): contact_info.append(f"Phone: {data.get('phone')}")
            if data.get('address'): contact_info.append(f"Address: {data.get('address')}")
            if data.get('linkedin'): contact_info.append(f"LinkedIn: {data.get('linkedin')}")
            if data.get('website'): contact_info.append(f"Website: {data.get('website')}")
            
            if contact_info:
                contact_para = doc.add_paragraph(' | '.join(contact_info))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_para_format = contact_para.paragraph_format
                contact_para_format.space_after = Pt(12)
            
            # Summary
            if data.get('summary'):
                doc.add_heading("Mục tiêu nghề nghiệp / Tóm tắt", level=1)
                summary_para = doc.add_paragraph(data.get('summary'))
                summary_para_format = summary_para.paragraph_format
                summary_para_format.space_after = Pt(12)

            # Kinh nghiệm
            experiences = processed_experiences
            if experiences and len(experiences) > 0:
                doc.add_heading("Kinh nghiệm làm việc", level=1)
                for exp in experiences:
                    # Vị trí và công ty
                    exp_title = doc.add_paragraph()
                    exp_title.add_run(f"{exp.get('position', '')}").bold = True
                    exp_title.add_run(f" - {exp.get('company', '')}")
                    if exp.get('start_date') or exp.get('end_date'):
                        dates = []
                        if exp.get('start_date'):
                            dates.append(exp.get('start_date'))
                        if exp.get('end_date'):
                            dates.append(exp.get('end_date'))
                        elif exp.get('is_current'):
                            dates.append('Hiện tại')
                        if dates:
                            exp_title.add_run(f" ({' - '.join(dates)})")
                    exp_title_format = exp_title.paragraph_format
                    exp_title_format.space_after = Pt(6)
                    
                    # Mô tả
                    if exp.get('description'):
                        desc_para = doc.add_paragraph(exp.get('description'))
                        desc_para_format = desc_para.paragraph_format
                        desc_para_format.left_indent = Inches(0.25)
                        desc_para_format.space_after = Pt(12)

            # Giáo dục
            education = processed_education
            if education and len(education) > 0:
                doc.add_heading("Học vấn", level=1)
                for edu in education:
                    edu_para = doc.add_paragraph()
                    if edu.get('school'):
                        edu_para.add_run(edu.get('school')).bold = True
                    if edu.get('major'):
                        edu_para.add_run(f" - {edu.get('major')}")
                    if edu.get('degree'):
                        edu_para.add_run(f" ({edu.get('degree')})")
                    if edu.get('year'):
                        edu_para.add_run(f" - {edu.get('year')}")
                    edu_para_format = edu_para.paragraph_format
                    edu_para_format.space_after = Pt(12)

            # Kỹ năng
            skills = context.get('skills', [])
            if skills and len(skills) > 0:
                doc.add_heading("Kỹ năng", level=1)
                skills_text = ', '.join(skills) if isinstance(skills, list) else str(skills)
                skills_para = doc.add_paragraph(skills_text)
                skills_para_format = skills_para.paragraph_format
                skills_para_format.space_after = Pt(12)

            # Chứng chỉ
            certifications = processed_certifications
            if certifications and len(certifications) > 0:
                doc.add_heading("Chứng chỉ / Giải thưởng", level=1)
                for cert in certifications:
                    cert_para = doc.add_paragraph()
                    if cert.get('name'):
                        cert_para.add_run(cert.get('name')).bold = True
                    if cert.get('organization'):
                        cert_para.add_run(f" - {cert.get('organization')}")
                    if cert.get('date'):
                        cert_para.add_run(f" ({cert.get('date')})")
                    cert_para_format = cert_para.paragraph_format
                    cert_para_format.space_after = Pt(12)

            # Dự án
            projects = processed_projects
            if projects and len(projects) > 0:
                doc.add_heading("Dự án", level=1)
                for proj in projects:
                    proj_para = doc.add_paragraph()
                    if proj.get('name'):
                        proj_para.add_run(proj.get('name')).bold = True
                    if proj.get('url'):
                        proj_para.add_run(f" - {proj.get('url')}")
                    proj_para_format = proj_para.paragraph_format
                    proj_para_format.space_after = Pt(6)
                    if proj.get('description'):
                        desc_para = doc.add_paragraph(proj.get('description'))
                        desc_para_format = desc_para.paragraph_format
                        desc_para_format.left_indent = Inches(0.25)
                        desc_para_format.space_after = Pt(12)

            # Ngôn ngữ
            languages = processed_languages
            if languages and len(languages) > 0:
                doc.add_heading("Ngôn ngữ", level=1)
                for lang in languages:
                    lang_text = f"{lang.get('name', '')} - {lang.get('level', '')}"
                    lang_para = doc.add_paragraph(lang_text)
                    lang_para_format = lang_para.paragraph_format
                    lang_para_format.space_after = Pt(12)

            f = io.BytesIO()
            doc.save(f)
            f.seek(0)
            return (
                f.read(),
                200,
                {
                    "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "Content-Disposition": f"attachment; filename={data.get('full_name','CV')}.docx",
                },
            )

        else:
            return jsonify({'success': False, 'message': 'Định dạng không hỗ trợ'}), 400

    except Exception as e:
        logger.error(f"Error export temp CV: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'message': f'Lỗi khi export CV: {str(e)}'}), 500


@cv_builder_bp.route('/api/cv-builder/upload-avatar', methods=['POST'])
@jwt_required()
def upload_avatar():
    """Upload ảnh đại diện cho CV"""
    try:
        user_id = get_user_id_from_jwt()
        
        if 'avatar' not in request.files:
            return jsonify({'success': False, 'message': 'Không có file ảnh'}), 400
        
        file = request.files['avatar']
        if file.filename == '':
            return jsonify({'success': False, 'message': 'Chưa chọn file'}), 400
        
        avatar_path = save_avatar(file, user_id)
        if avatar_path:
            return jsonify({
                'success': True,
                'avatar_url': avatar_path
            }), 200
        else:
            return jsonify({'success': False, 'message': 'Định dạng file không hợp lệ'}), 400
    
    except Exception as e:
        logger.error(f"Error uploading avatar: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'message': f'Lỗi khi upload ảnh: {str(e)}'}), 500

@cv_builder_bp.route('/api/cv-builder/preview-temp', methods=['POST'])
@jwt_required()
def preview_temp_cv():
    try:
        if request.is_json:
            data = request.get_json() or {}
        else:
            json_data = request.form.get('data')
            if json_data:
                data = json.loads(json_data)
            else:
                data = {}
        
        template = data.get('template', 'classic')
        template_file = f"cv_formats/{template}.html"

        experiences = data.get('experiences', [])
        processed_experiences = []
        for exp in experiences:
            processed_exp = {
                'position': exp.get('position', ''),
                'company': exp.get('company', ''),
                'start_date': exp.get('start_date', ''),
                'end_date': exp.get('end_date', ''),
                'is_current': exp.get('is_current', False),
                'description': exp.get('description', '')
            }
            if processed_exp['end_date'] == '' and processed_exp['is_current']:
                processed_exp['end_date'] = None
            processed_experiences.append(processed_exp)

        education = data.get('education', [])
        processed_education = []
        for edu in education:
            processed_edu = {
                'school': edu.get('school', ''),
                'major': edu.get('major', ''),
                'degree': edu.get('degree', ''),
                'year': edu.get('year', '')
            }
            processed_education.append(processed_edu)

        skills = data.get('skills', [])
        logger.info(f"Preview CV - Raw skills from data: {repr(skills)}, type: {type(skills)}")
        
        if isinstance(skills, str):
            skills = [s.strip() for s in skills.replace('\n', ',').split(',') if s.strip()]
        elif isinstance(skills, list):
            skills = [str(s).strip() for s in skills if s and str(s).strip()]
        else:
            skills = []
        
        logger.info(f"Preview CV - Processed skills: {repr(skills)}, count: {len(skills)}")

        certifications = data.get('certifications', [])
        processed_certifications = []
        for cert in certifications:
            processed_cert = {
                'name': cert.get('name', ''),
                'organization': cert.get('organization', ''),
                'date': cert.get('date', '')
            }
            processed_certifications.append(processed_cert)

        projects = data.get('projects', [])
        processed_projects = []
        for proj in projects:
            processed_proj = {
                'name': proj.get('name', ''),
                'url': proj.get('url', ''),
                'description': proj.get('description', '')
            }
            processed_projects.append(processed_proj)

        languages = data.get('languages', [])
        processed_languages = []
        for lang in languages:
            processed_lang = {
                'name': lang.get('name', ''),
                'level': lang.get('level', '')
            }
            processed_languages.append(processed_lang)

        avatar_url = data.get('avatar_url', '')
        if 'avatar' in request.files:
            user_id = get_user_id_from_jwt()
            file = request.files['avatar']
            if file and file.filename != '':
                avatar_path = save_avatar(file, user_id)
                if avatar_path:
                    avatar_url = avatar_path
        
        try:
            from app.routes.recruiter_routes import safe_decode_text, fix_common_vietnamese_errors
        except ImportError:
            def safe_decode_text(text):
                if isinstance(text, bytes):
                    return text.decode('utf-8', errors='replace')
                return str(text) if text else ''
            def fix_common_vietnamese_errors(text):
                return text
        
        def ensure_utf8(text):
            if not text:
                return text
            text = safe_decode_text(text)
            text = fix_common_vietnamese_errors(text)
            return text
        
        full_name = ensure_utf8(data.get('full_name', ''))
        email = ensure_utf8(data.get('email', ''))
        phone = ensure_utf8(data.get('phone', ''))
        address = ensure_utf8(data.get('address', ''))
        linkedin = ensure_utf8(data.get('linkedin', ''))
        website = ensure_utf8(data.get('website', ''))
        summary = ensure_utf8(data.get('summary', ''))
        
        for exp in processed_experiences:
            exp['position'] = ensure_utf8(exp.get('position', ''))
            exp['company'] = ensure_utf8(exp.get('company', ''))
            exp['description'] = ensure_utf8(exp.get('description', ''))
        
        for edu in processed_education:
            edu['school'] = ensure_utf8(edu.get('school', ''))
            edu['major'] = ensure_utf8(edu.get('major', ''))
            edu['degree'] = ensure_utf8(edu.get('degree', ''))
        
        if isinstance(skills, list):
            skills = [ensure_utf8(str(skill).strip()) for skill in skills if skill and str(skill).strip()]
        else:
            skills = []
        
        for cert in processed_certifications:
            cert['name'] = ensure_utf8(cert.get('name', ''))
            cert['organization'] = ensure_utf8(cert.get('organization', ''))
        
        for proj in processed_projects:
            proj['name'] = ensure_utf8(proj.get('name', ''))
            proj['description'] = ensure_utf8(proj.get('description', ''))
        
        for lang in processed_languages:
            lang['name'] = ensure_utf8(lang.get('name', ''))
            lang['level'] = ensure_utf8(lang.get('level', ''))
        
        logger.info(f"Preview CV - Skills received: {repr(skills)}")
        logger.info(f"Preview CV - Skills count: {len(skills) if isinstance(skills, list) else 0}")
        
        logger.info(f"Preview CV - Summary received: {repr(summary)}")
        logger.info(f"Preview CV - Summary after ensure_utf8: {repr(summary)}")
        
        final_skills = skills if isinstance(skills, list) and len(skills) > 0 else []
        
        context = {
            'full_name': full_name or '',
            'email': email or '',
            'phone': phone or '',
            'address': address or '',
            'linkedin': linkedin or '',
            'website': website or '',
            'summary': summary or '',
            'avatar_url': avatar_url or '',
            'experiences': processed_experiences or [],
            'education': processed_education or [],
            'skills': final_skills,
            'certifications': processed_certifications or [],
            'projects': processed_projects or [],
            'languages': processed_languages or []
        }
        
        # Debug: Log context để kiểm tra
        logger.info(f"Preview CV - Context summary: {repr(context.get('summary'))}")
        logger.info(f"Preview CV - Context skills: {repr(context.get('skills'))}")
        logger.info(f"Preview CV - Context skills count: {len(context.get('skills', []))}")
        logger.info(f"Preview CV - Context skills type: {type(context.get('skills'))}")

        return render_template(template_file, **context)

    except Exception as e:
        logger.error(f"Error preview temp CV: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'message': str(e)}), 500



@cv_builder_bp.route('/api/cv-builder/enhance', methods=['POST'])
@jwt_required()
def enhance_cv_with_ai():
    """API sử dụng AI để cải thiện nội dung CV"""
    try:
        data = request.get_json()
        enhancement_type = data.get('type') 
        
        if enhancement_type == 'summary':
            summary = data.get('summary')
            if not summary:
                return jsonify({'success': False, 'message': 'Vui lòng nhập nội dung tóm tắt'}), 400
            
            enhanced = enhance_summary_with_ai(summary, data.get('job_title'))
            if enhanced is None:
                return jsonify({
                    'success': False,
                    'message': 'Không thể cải thiện nội dung. Vui lòng kiểm tra OPENAI_API_KEY trong file .env hoặc thử lại sau.'
                }), 500
            
            return jsonify({
                'success': True,
                'enhanced': enhanced,
                'message': 'Đã cải thiện nội dung thành công!'
            }), 200
        
        elif enhancement_type == 'experience':
            experience = data.get('experience')
            if not experience:
                return jsonify({'success': False, 'message': 'Vui lòng nhập thông tin kinh nghiệm'}), 400
            
            enhanced = enhance_experience_with_ai(experience)
            if enhanced is None or enhanced == experience:
                return jsonify({
                    'success': False,
                    'message': 'Không thể cải thiện nội dung. Vui lòng kiểm tra OPENAI_API_KEY trong file .env hoặc thử lại sau.'
                }), 500
            
            return jsonify({
                'success': True,
                'enhanced': enhanced,
                'message': 'Đã cải thiện kinh nghiệm thành công!'
            }), 200
        
        elif enhancement_type == 'skills':
            skills = data.get('skills', [])
            experiences = data.get('experiences', [])
            
            if not skills:
                return jsonify({'success': False, 'message': 'Vui lòng nhập danh sách kỹ năng'}), 400
            
            enhanced = enhance_skills_with_ai(skills, experiences)
            if not enhanced or enhanced == skills:
                return jsonify({
                    'success': False,
                    'message': 'Không thể cải thiện kỹ năng. Vui lòng kiểm tra OPENAI_API_KEY trong file .env hoặc thử lại sau.'
                }), 500
            
            return jsonify({
                'success': True,
                'enhanced': enhanced,
                'message': 'Đã tối ưu hóa kỹ năng thành công!'
            }), 200
        
        elif enhancement_type == 'full':
            cv_data = data.get('cv_data')
            if not cv_data:
                return jsonify({'success': False, 'message': 'Vui lòng cung cấp dữ liệu CV'}), 400
            
            enhanced = enhance_full_cv_with_ai(cv_data)
            if not enhanced or enhanced == cv_data:
                return jsonify({
                    'success': False,
                    'message': 'Không thể cải thiện CV. Vui lòng kiểm tra OPENAI_API_KEY trong file .env hoặc thử lại sau.'
                }), 500
            
            return jsonify({
                'success': True,
                'enhanced': enhanced,
                'message': 'Đã cải thiện toàn bộ CV thành công!'
            }), 200
        
        else:
            return jsonify({'success': False, 'message': 'Loại cải thiện không hợp lệ'}), 400
    
    except Exception as e:
        logger.error(f"Error enhancing CV with AI: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Lỗi khi cải thiện CV: {str(e)}'
        }), 500

@cv_builder_bp.route('/api/cv-builder/list', methods=['GET'])
@jwt_required()
def list_cv_data():
    """API lấy danh sách CV đã tạo"""
    try:
        user_id = get_user_id_from_jwt()
        
        cv_list = CVData.query.filter_by(user_id=user_id, is_active=True).order_by(CVData.updated_at.desc()).all()
        
        return jsonify({
            'success': True,
            'cvs': [{
                'id': cv.id,
                'full_name': cv.full_name,
                'template': cv.template,
                'category': cv.category.name if cv.category else None,
                'created_at': cv.created_at.isoformat() if cv.created_at else None,
                'updated_at': cv.updated_at.isoformat() if cv.updated_at else None
            } for cv in cv_list]
        }), 200
    
    except Exception as e:
        logger.error(f"Error listing CV data: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Lỗi khi tải danh sách CV: {str(e)}'
        }), 500

@cv_builder_bp.route('/api/cv-builder/<int:cv_id>', methods=['GET'])
@jwt_required()
def get_cv_data(cv_id):
    """API lấy chi tiết CV"""
    try:
        user_id = get_user_id_from_jwt()
        
        cv_data = CVData.query.get_or_404(cv_id)
        
        if cv_data.user_id != user_id:
            return jsonify({'success': False, 'message': 'Không có quyền truy cập CV này'}), 403
        
        return jsonify({
            'success': True,
            'cv': cv_data.to_dict()
        }), 200
    
    except Exception as e:
        logger.error(f"Error getting CV data: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Lỗi khi tải CV: {str(e)}'
        }), 500

@cv_builder_bp.route('/api/cv-builder/<int:cv_id>', methods=['DELETE'])
@jwt_required()
def delete_cv_data(cv_id):
    """API xóa CV"""
    try:
        user_id = get_user_id_from_jwt()
        
        cv_data = CVData.query.get_or_404(cv_id)
        
        if cv_data.user_id != user_id:
            return jsonify({'success': False, 'message': 'Không có quyền xóa CV này'}), 403
        
        cv_data.is_active = False
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Xóa CV thành công'
        }), 200
    
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting CV data: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Lỗi khi xóa CV: {str(e)}'
        }), 500

@cv_builder_bp.route('/cv/<int:cv_id>/preview', methods=['GET'])
@jwt_required()
def preview_cv(cv_id):
    """Trang preview CV"""
    return render_template('cv_preview.html', cv_id=cv_id)

@cv_builder_bp.route('/cv/<int:cv_id>/export', methods=['GET'])
@jwt_required()
def export_cv(cv_id):
    """Export CV ra PDF/DOCX"""
    try:
        user_id = get_user_id_from_jwt()
        cv_data = CVData.query.get_or_404(cv_id)
        
        if cv_data.user_id != user_id:
            return jsonify({'success': False, 'message': 'Không có quyền export CV này'}), 403
        
        format_type = request.args.get('format', 'pdf')  # pdf or docx
        template = cv_data.template or 'classic'
        
        # Render template
        cv_dict = cv_data.to_dict()
        # Use AI enhanced content if available
        if cv_data.ai_enhanced_summary:
            cv_dict['summary'] = cv_data.ai_enhanced_summary
        if cv_data.ai_enhanced_experiences:
            cv_dict['experiences'] = cv_data.ai_enhanced_experiences
        if cv_data.ai_enhanced_skills:
            cv_dict['skills'] = cv_data.ai_enhanced_skills
        
        if format_type == 'pdf':
            # TODO: Implement PDF export using pdfkit or weasyprint
            return jsonify({'success': False, 'message': 'PDF export chưa được triển khai'}), 501
        elif format_type == 'docx':
            # TODO: Implement DOCX export using python-docx
            return jsonify({'success': False, 'message': 'DOCX export chưa được triển khai'}), 501
        else:
            return jsonify({'success': False, 'message': 'Định dạng không hỗ trợ'}), 400
    
    except Exception as e:
        logger.error(f"Error exporting CV: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Lỗi khi export CV: {str(e)}'
        }), 500

