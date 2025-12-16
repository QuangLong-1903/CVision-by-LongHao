import os
import logging
from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify
from werkzeug.utils import secure_filename
from app.extensions import db
from app.models import CV, User, JobCategory, ClassificationLog
from app.models.cvprocessingqueue import CVProcessingQueue
from flask_jwt_extended import jwt_required, get_jwt_identity, decode_token
from flask_jwt_extended.exceptions import JWTDecodeError
from app.utils.text_extractor import extract_text_from_file
from app.utils.classifier import classify_cv_by_keywords, get_category_id_by_name
from datetime import datetime, timedelta
from reportlab.pdfgen import canvas
from sqlalchemy import func
logger = logging.getLogger(__name__)

cv_bp = Blueprint('cv', __name__)

# Đường dẫn tuyệt đối đến thư mục uploads
# Dựa trên vị trí file hiện tại (cv_routes.py)
# cv_routes.py nằm ở: Flask_CVProject/app/routes/cv_routes.py
# Cần lấy: Flask_CVProject/app/static/uploads
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # Flask_CVProject/app
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'static', 'uploads')
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'jpg', 'jpeg', 'png'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@cv_bp.route("/api/cv/classify", methods=["POST"])
def classify_cv():
    data = request.get_json()
    cv_text = data.get("cv_text")

    category, confidence = classify_cv_by_keywords(cv_text)
    return jsonify({"category": category, "confidence": confidence})

from flask import Blueprint, send_file
from app.models.cv import CV
import io
from reportlab.pdfgen import canvas

from flask import jsonify, render_template
from app.models.cv import CV
from docx import Document
from io import BytesIO
from xhtml2pdf import pisa

@cv_bp.route("/api/cv/<int:cv_id>/export", methods=["GET"])
def export_cv(cv_id):
    cv = CV.query.get_or_404(cv_id)
    export_format = request.args.get("format", "pdf")
    template = request.args.get("template", "classic")
    print("✅ Đang dùng phiên bản xhtml2pdf")


    rendered_html = render_template(
        f"cv_formats/{template}.html",
        full_name=cv.full_name,
        email=cv.email,
        phone=cv.phone,
        address=cv.address,
        summary=cv.summary,
        experiences=cv.experiences,
        education=cv.education,
        skills=cv.skills,
        certifications=cv.certifications,
        activities=cv.activities,
        projects=cv.projects,
        languages=cv.languages
    )

    filename = f"{cv.full_name.replace(' ', '_')}_CV.{export_format}"

    if export_format == "pdf":
        try:
            pdf_stream = BytesIO()
            pisa_status = pisa.CreatePDF(rendered_html, dest=pdf_stream)
            if pisa_status.err:
                return jsonify({"error": "Lỗi khi export PDF bằng xhtml2pdf"}), 500
            pdf_stream.seek(0)
            return (
                pdf_stream.read(),
                200,
                {
                    "Content-Type": "application/pdf",
                    "Content-Disposition": f"attachment; filename={filename}",
                },
            )
        except Exception as e:
            return jsonify({"error": f"Lỗi khi export PDF: {str(e)}"}), 500

    elif export_format == "docx":
        doc = Document()
        doc.add_heading(cv.full_name, 0)
        doc.add_paragraph(f"Email: {cv.email} | Phone: {cv.phone} | Address: {cv.address}")
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(cv.summary)

        doc.add_heading("Experience", level=1)
        for exp in cv.experiences:
            doc.add_paragraph(
                f"{exp['position']} - {exp['company']} "
                f"({exp.get('start_date', '')} - {exp.get('end_date', '')})\n{exp['description']}"
            )

        doc.add_heading("Education", level=1)
        for edu in cv.education:
            doc.add_paragraph(f"{edu['school']} - {edu['major']} ({edu['degree']} {edu['year']})")

        doc.add_heading("Skills", level=1)
        for skill in cv.skills:
            doc.add_paragraph(skill)

        doc.add_heading("Certifications", level=1)
        for cert in cv.certifications:
            doc.add_paragraph(f"{cert['name']} - {cert['organization']} ({cert['date']})")

        doc.add_heading("Projects", level=1)
        for proj in cv.projects:
            doc.add_paragraph(f"{proj['name']} - {proj['url']}\n{proj['description']}")

        doc.add_heading("Languages", level=1)
        for lang in cv.languages:
            doc.add_paragraph(f"{lang['name']} - {lang['level']}")

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return (
            file_stream.read(),
            200,
            {
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "Content-Disposition": f"attachment; filename={filename}",
            },
        )

    return jsonify({"error": "Unsupported format"}), 400



@cv_bp.route("/api/cv/top", methods=["GET"])
def top_cvs():
    """
    Lấy danh sách top CV ưu tú theo ngành nghề
    """
    category = request.args.get("category")
    limit = int(request.args.get("limit", 10))

    # Query DB: lọc CV theo category, sắp xếp theo tiêu chí đánh giá
    cvs = CV.query.join(JobCategory).filter(JobCategory.name == category).limit(limit).all()

    return jsonify([
        {
            "id": cv.id,
            "file_name": cv.file_name,
            "user_id": cv.user_id,
            "uploaded_at": cv.uploaded_at.isoformat() if cv.uploaded_at else None
        }
        for cv in cvs
    ])



def get_user_id_from_jwt():
    """Helper function to get user_id from JWT identity and convert to int"""
    user_id = get_jwt_identity()
    # Convert string to int if needed (JWT subject must be string, but DB uses int)
    if isinstance(user_id, str) and user_id.isdigit():
        return int(user_id)
    return user_id

def get_user_id_from_request():
    """Extract user_id from JWT token if present (optional)"""
    try:
        auth_header = request.headers.get('Authorization', '')
        if auth_header and auth_header.startswith('Bearer '):
            token = auth_header.split(' ')[1]
            decoded = decode_token(token)
            user_id = decoded.get('sub')  # JWT identity (user_id) - should be string
            if user_id:
                # Convert string to int for database query
                user_id = int(user_id) if isinstance(user_id, str) and user_id.isdigit() else user_id
                logger.info(f"Extracted user_id from token: {user_id}")
            return user_id
    except (JWTDecodeError, Exception) as e:
        logger.warning(f"Error extracting user_id from token: {str(e)}")
        pass
    return None

@cv_bp.route('/', methods=['GET'])
def index():
    """Trang chủ"""
    return render_template('index.html')

@cv_bp.route('/create-cv', methods=['GET'])
def create_cv():
    """Trang tạo CV thông minh"""
    return render_template('create_cv.html')

@cv_bp.route('/upload', methods=['GET', 'POST'])
def upload_cv():
    """Redirect to create-cv (legacy support)"""
    if request.method == 'GET':
        return redirect(url_for('cv.create_cv'))
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('Không tìm thấy file tải lên.', 'danger')
            return redirect(request.url)

        file = request.files['file']
        if file.filename == '':
            flash('Chưa chọn file.', 'warning')
            return redirect(request.url)

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            os.makedirs(UPLOAD_FOLDER, exist_ok=True)
            
            # Lưu file và lấy kích thước
            file.save(file_path)
            file_size = os.path.getsize(file_path)
            file_type = filename.rsplit('.', 1)[1].lower()

            # Extract text from file
            # For images, skip OCR if file is too large or processing takes too long
            extracted_text = None
            try:
                # Check file size before OCR for images
                if file_type in ['jpg', 'jpeg', 'png']:
                    file_size_mb = file_size / (1024 * 1024)
                    if file_size_mb > 5:
                        logger.info(f"Image {filename} is large ({file_size_mb:.1f}MB). Skipping OCR to prevent hanging. File uploaded successfully.")
                        extracted_text = None
                    else:
                        logger.info(f"Processing image {filename} with OCR (may take 10-30 seconds)...")
                        try:
                            extracted_text = extract_text_from_file(file_path, file_type)
                        except Exception as ocr_error:
                            logger.warning(f"OCR failed for {filename}: {str(ocr_error)}. File uploaded but text extraction skipped.")
                            extracted_text = None
                else:
                    extracted_text = extract_text_from_file(file_path, file_type)
                
                if extracted_text:
                    logger.info(f"Successfully extracted text from {filename} ({len(extracted_text)} characters)")
                else:
                    if file_type in ['jpg', 'jpeg', 'png']:
                        logger.info(f"Image {filename} uploaded successfully. Text extraction skipped (this is normal for large images).")
                    else:
                        logger.warning(f"No text extracted from {filename}")
            except Exception as e:
                logger.error(f"Error extracting text from {filename}: {str(e)}")
                # Continue even if extraction fails - file is still saved

            # Lấy user_id nếu user đã login (từ JWT token nếu có)
            user_id = get_user_id_from_request()

            # Classify CV if text was extracted
            predicted_category_id = None
            confidence = None
            category_name = None
            if extracted_text:
                try:
                    category_name, confidence_score = classify_cv_by_keywords(extracted_text)
                    if category_name:
                        predicted_category_id = get_category_id_by_name(category_name, JobCategory)
                        if predicted_category_id:
                            confidence = confidence_score
                            logger.info(f"Classified CV as '{category_name}' (ID: {predicted_category_id}) with confidence {confidence_score:.2%}")
                        else:
                            logger.warning(f"Category '{category_name}' not found in database")
                except Exception as e:
                    logger.error(f"Error classifying CV: {str(e)}")
                    # Continue even if classification fails

            new_cv = CV(
                file_name=filename,
                file_type=file_type,
                file_size=file_size,
                user_id=user_id,
                file_content=extracted_text,
                predicted_category_id=predicted_category_id
            )
            db.session.add(new_cv)
            db.session.commit()
            
            # Create classification log if classification was successful
            if predicted_category_id and confidence is not None and user_id:
                try:
                    classification_log = ClassificationLog(
                        cv_id=new_cv.id,
                        predicted_category_id=predicted_category_id,
                        confidence=confidence,
                        user_id=user_id
                    )
                    db.session.add(classification_log)
                    db.session.commit()
                    logger.info(f"Created classification log for CV {new_cv.id}")
                except Exception as e:
                    logger.error(f"Error creating classification log: {str(e)}")
                    db.session.rollback()

            # Prepare success message with classification result
            if predicted_category_id and category_name and confidence is not None:
                flash(f'Tải lên thành công! CV đã được phân loại: {category_name} (Confidence: {confidence:.1%})', 'success')
            else:
                flash('Tải lên thành công!', 'success')
            
            # Redirect to create-cv page if user is logged in, otherwise stay on upload page
            if user_id:
                return redirect(url_for('cv.create_cv'))
            else:
                return redirect(url_for('cv.upload_cv'))
        else:
            flash('Định dạng file không được hỗ trợ. Chỉ hỗ trợ: PDF, DOCX, TXT, JPG, JPEG, PNG', 'danger')
            return redirect(request.url)

    return render_template('upload_cv.html')


@cv_bp.route('/about', methods=['GET'])
def about():
    """Trang giới thiệu"""
    return render_template('about.html')


@cv_bp.route('/jobs', methods=['GET'])
def jobs():
    """Trang danh sách jobs cho ứng viên"""
    return render_template('jobs.html')

@cv_bp.route('/jobs/<int:job_id>', methods=['GET'])
def job_detail(job_id):
    """Trang chi tiết job cho ứng viên"""
    return render_template('job_detail.html', job_id=job_id)


# Route my-cvs đã bị xóa - không còn cần thiết


@cv_bp.route('/cv/<int:cv_id>', methods=['GET'])
def cv_detail(cv_id):
    """Trang chi tiết CV"""
    return render_template('cv_detail.html', cv_id=cv_id)


@cv_bp.route('/api/categories', methods=['GET'])
@jwt_required()
def get_categories():
    """API lấy danh sách categories để filter"""
    try:
        categories = JobCategory.query.order_by(JobCategory.name).all()
        
        category_list = []
        for category in categories:
            category_list.append({
                'id': category.id,
                'name': category.name,
                'description': category.description
            })
        
        return jsonify({
            'success': True,
            'categories': category_list
        }), 200
    
    except Exception as e:
        logger.error(f"Error in get_categories endpoint: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Error loading categories: {str(e)}'
        }), 500


@cv_bp.route('/api/statistics', methods=['GET'])
@jwt_required()
def get_statistics():
    """
    API lấy statistics về CVs của user
    Returns: Total CVs, CVs by Category, Average Confidence, Upload Trends
    """
    try:
        current_user_id = get_user_id_from_jwt()
        
        # Get user's CVs
        user_cvs = CV.query.filter_by(user_id=current_user_id).all()
        total_cvs = len(user_cvs)
        
        # Calculate statistics
        stats = {
            'total_cvs': total_cvs,
            'average_confidence': 0.0,
            'cvs_by_category': {},
            'cvs_by_file_type': {},
            'upload_trends': [],
            'most_common_categories': []
        }
        
        if total_cvs == 0:
            return jsonify({
                'success': True,
                'statistics': stats
            }), 200
        
        # Get classification logs for user's CVs
        cv_ids = [cv.id for cv in user_cvs]
        logs = ClassificationLog.query.filter(ClassificationLog.cv_id.in_(cv_ids)).all()
        
        # Calculate average confidence
        total_confidence = 0.0
        confidence_count = 0
        category_counts = {}
        file_type_counts = {}
        
        # CVs by Category
        for cv in user_cvs:
            # Count by category
            if cv.predicted_category_id:
                category = JobCategory.query.get(cv.predicted_category_id)
                if category:
                    category_name = category.name
                    category_counts[category_name] = category_counts.get(category_name, 0) + 1
            
            # Count by file type
            file_type = cv.file_type or 'unknown'
            file_type_counts[file_type] = file_type_counts.get(file_type, 0) + 1
        
        # Average confidence from logs
        for log in logs:
            if log.confidence is not None:
                total_confidence += log.confidence
                confidence_count += 1
        
        if confidence_count > 0:
            stats['average_confidence'] = round((total_confidence / confidence_count) * 100, 2)
        
        # CVs by Category (format for charts)
        stats['cvs_by_category'] = [
            {'name': name, 'count': count}
            for name, count in sorted(category_counts.items(), key=lambda x: x[1], reverse=True)
        ]
        
        # CVs by File Type
        stats['cvs_by_file_type'] = [
            {'type': file_type.upper(), 'count': count}
            for file_type, count in sorted(file_type_counts.items(), key=lambda x: x[1], reverse=True)
        ]
        
        # Upload trends (last 30 days)
        thirty_days_ago = datetime.utcnow() - timedelta(days=30)
        recent_cvs = CV.query.filter(
            CV.user_id == current_user_id,
            CV.uploaded_at >= thirty_days_ago
        ).order_by(CV.uploaded_at.asc()).all()
        
        # Group by date
        date_counts = {}
        for cv in recent_cvs:
            if cv.uploaded_at:
                date_str = cv.uploaded_at.strftime('%Y-%m-%d')
                date_counts[date_str] = date_counts.get(date_str, 0) + 1
        
        # Format for timeline chart
        stats['upload_trends'] = [
            {'date': date, 'count': count}
            for date, count in sorted(date_counts.items())
        ]
        
        # Most common categories (top 5)
        stats['most_common_categories'] = [
            {'name': name, 'count': count}
            for name, count in sorted(category_counts.items(), key=lambda x: x[1], reverse=True)[:5]
        ]
        
        return jsonify({
            'success': True,
            'statistics': stats
        }), 200
    
    except Exception as e:
        logger.error(f"Error in get_statistics endpoint: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Error loading statistics: {str(e)}'
        }), 500


@cv_bp.route('/api/cvs', methods=['GET'])
@jwt_required()
def get_cvs():
    """
    API lấy danh sách CV của user
    Hỗ trợ search, filter, và sort
    
    Query Parameters:
    - search: Tìm kiếm theo tên file (optional)
    - category: Filter theo category ID (optional)
    - file_type: Filter theo file type (pdf, docx, txt, jpg, jpeg, png) (optional)
    - sort_by: Sort field (name, date, size, category) (default: date)
    - sort_order: Sort order (asc, desc) (default: desc)
    """
    try:
        current_user_id = get_user_id_from_jwt()
        logger.info(f"Loading CVs for user_id: {current_user_id}")
        
        # Get query parameters
        search = request.args.get('search', '').strip()
        category_id = request.args.get('category', type=int)
        file_type = request.args.get('file_type', '').strip().lower()
        sort_by = request.args.get('sort_by', 'date').lower()
        sort_order = request.args.get('sort_order', 'desc').lower()
        
        # Start with base query
        query = CV.query.filter_by(user_id=current_user_id)
        
        # Apply search filter (search in file name)
        if search:
            query = query.filter(CV.file_name.ilike(f'%{search}%'))
        
        # Apply category filter
        if category_id:
            query = query.filter(CV.predicted_category_id == category_id)
        
        # Apply file type filter
        if file_type:
            # Normalize file type
            if file_type in ['jpg', 'jpeg']:
                query = query.filter(CV.file_type.in_(['jpg', 'jpeg']))
            else:
                query = query.filter(CV.file_type == file_type)
        
        # Apply sorting
        if sort_by == 'name':
            if sort_order == 'asc':
                query = query.order_by(CV.file_name.asc())
            else:
                query = query.order_by(CV.file_name.desc())
        elif sort_by == 'size':
            if sort_order == 'asc':
                query = query.order_by(CV.file_size.asc())
            else:
                query = query.order_by(CV.file_size.desc())
        elif sort_by == 'category':
            if sort_order == 'asc':
                query = query.order_by(CV.predicted_category_id.asc())
            else:
                query = query.order_by(CV.predicted_category_id.desc())
        else:  # Default: sort by date
            if sort_order == 'asc':
                query = query.order_by(CV.uploaded_at.asc())
            else:
                query = query.order_by(CV.uploaded_at.desc())
        
        # Get all CVs
        cvs = query.all()
        logger.info(f"Found {len(cvs)} CVs for user {current_user_id} (filters: search={search}, category={category_id}, file_type={file_type}, sort={sort_by} {sort_order})")
        
        # Format response
        cv_list = []
        for cv in cvs:
            try:
                # Safely get category name
                category_name = None
                if cv.predicted_category_id:
                    try:
                        # Try to get category directly from database if relationship fails
                        category = JobCategory.query.get(cv.predicted_category_id)
                        if category:
                            category_name = category.name
                        else:
                            # Fallback to relationship
                            try:
                                category_name = cv.category.name if cv.category else None
                            except:
                                category_name = None
                    except Exception as e:
                        logger.warning(f"Error getting category name for CV {cv.id}: {str(e)}")
                        category_name = None
                
                cv_data = {
                    'id': cv.id,
                    'file_name': cv.file_name,
                    'file_type': cv.file_type,
                    'file_size': cv.file_size,
                    'uploaded_at': cv.uploaded_at.isoformat() if cv.uploaded_at else None,
                    'predicted_category': category_name,
                    'predicted_category_id': cv.predicted_category_id
                }
                cv_list.append(cv_data)
            except Exception as e:
                logger.error(f"Error formatting CV {cv.id}: {str(e)}")
                # Continue with other CVs even if one fails
                continue
        
        return jsonify({
            'success': True,
            'cvs': cv_list,
            'total': len(cv_list),
            'filters': {
                'search': search,
                'category_id': category_id,
                'file_type': file_type,
                'sort_by': sort_by,
                'sort_order': sort_order
            }
        }), 200
    
    except Exception as e:
        logger.error(f"Error in get_cvs endpoint: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Error loading CVs: {str(e)}'
        }), 500


@cv_bp.route('/api/cvs/<int:cv_id>', methods=['GET'])
@jwt_required()
def get_cv_detail(cv_id):
    """API lấy chi tiết CV của user hoặc cho recruiter xem CV ứng viên"""
    try:
        current_user_id = get_user_id_from_jwt()
        user = User.query.get(current_user_id)
        
        # Tìm CV
        cv = CV.query.get_or_404(cv_id)
        
        # Kiểm tra quyền: user có thể xem CV của mình, hoặc recruiter có thể xem CV của ứng viên đã apply vào job của họ
        can_view = False
        if cv.user_id == current_user_id:
            can_view = True
        elif user.role == 'recruiter':
            # Kiểm tra xem CV này có thuộc về ứng viên đã apply vào job của recruiter không
            from app.models import JobApplication, JobPosting
            application = JobApplication.query.join(JobPosting).filter(
                JobApplication.cv_id == cv_id,
                JobPosting.recruiter_id == current_user_id
            ).first()
            can_view = application is not None
        
        if not can_view:
            return jsonify({
                'success': False,
                'message': 'You do not have permission to view this CV'
            }), 403
        
        # Format response
        category_name = cv.category.name if cv.category else None
        
        # Get classification logs
        classification_logs = []
        if cv.logs:
            for log in cv.logs:
                classification_logs.append({
                    'id': log.id,
                    'predicted_category': log.category.name if log.category else None,
                    'predicted_category_id': log.predicted_category_id,
                    'confidence': log.confidence,
                    'created_at': log.created_at.isoformat() if log.created_at else None
                })
        
        cv_data = {
            'id': cv.id,
            'file_name': cv.file_name,
            'file_type': cv.file_type,
            'file_size': cv.file_size,
            'uploaded_at': cv.uploaded_at.isoformat() if cv.uploaded_at else None,
            'predicted_category': category_name,
            'predicted_category_id': cv.predicted_category_id,
            'file_content': cv.file_content,  # Extracted text
            'classification_logs': classification_logs
        }
        
        return jsonify({
            'success': True,
            'cv': cv_data
        }), 200
    
    except Exception as e:
        logger.error(f"Error loading CV detail: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Error loading CV: {str(e)}'
        }), 500

@cv_bp.route('/api/cvs/<int:cv_id>/download', methods=['GET'])
@jwt_required()
def download_cv(cv_id):
    """API download CV file - cho phép user download CV của mình hoặc recruiter download CV ứng viên"""
    try:
        from flask import send_file
        current_user_id = get_user_id_from_jwt()
        user = User.query.get(current_user_id)
        
        # Tìm CV
        cv = CV.query.get_or_404(cv_id)
        
        # Kiểm tra quyền: user có thể download CV của mình, hoặc recruiter có thể download CV của ứng viên đã apply
        can_download = False
        if cv.user_id == current_user_id:
            can_download = True
        elif user.role == 'recruiter':
            # Kiểm tra xem CV này có thuộc về ứng viên đã apply vào job của recruiter không
            from app.models import JobApplication, JobPosting
            application = JobApplication.query.join(JobPosting).filter(
                JobApplication.cv_id == cv_id,
                JobPosting.recruiter_id == current_user_id
            ).first()
            can_download = application is not None
        
        if not can_download:
            return jsonify({
                'success': False,
                'message': 'You do not have permission to download this CV'
            }), 403
        
        # Tìm file path (đảm bảo đường dẫn tuyệt đối)
        file_path = os.path.join(UPLOAD_FOLDER, cv.file_name)
        file_path = os.path.abspath(file_path)  # Chuyển thành đường dẫn tuyệt đối
        
        # Log để debug
        logger.info(f"Looking for CV file at: {file_path}")
        logger.info(f"File exists: {os.path.exists(file_path)}")
        
        if not os.path.exists(file_path):
            # Thử tìm trong các vị trí khác có thể
            alternative_paths = [
                os.path.join(BASE_DIR, 'static', 'uploads', cv.file_name),  # Flask_CVProject/app/static/uploads
                os.path.join(os.getcwd(), 'Flask_CVProject', 'app', 'static', 'uploads', cv.file_name),
                os.path.join('Flask_CVProject', 'app', 'static', 'uploads', cv.file_name)
            ]
            
            found = False
            for alt_path in alternative_paths:
                alt_path = os.path.abspath(alt_path)
                if os.path.exists(alt_path):
                    file_path = alt_path
                    found = True
                    logger.info(f"Found CV file at alternative path: {file_path}")
                    break
            
            if not found:
                logger.error(f"CV file not found at any path. Tried: {file_path} and alternatives")
                return jsonify({
                    'success': False,
                    'message': f'CV file not found: {cv.file_name}'
                }), 404
        
        return send_file(
            file_path,
            as_attachment=True,
            download_name=cv.file_name
        )
    
    except Exception as e:
        logger.error(f"Error downloading CV: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Error downloading CV: {str(e)}'
        }), 500

@cv_bp.route('/api/cvs/<int:cv_id>', methods=['DELETE'])
@jwt_required()
def delete_cv(cv_id):
    """API xóa CV của user"""
    try:
        current_user_id = get_user_id_from_jwt()
        
        # Tìm CV và kiểm tra ownership
        cv = CV.query.get_or_404(cv_id)
        
        if cv.user_id != current_user_id:
            return jsonify({
                'success': False,
                'message': 'You do not have permission to delete this CV'
            }), 403
        
        # Xóa tất cả ClassificationLog liên quan trước
        classification_logs = ClassificationLog.query.filter_by(cv_id=cv_id).all()
        if classification_logs:
            for log in classification_logs:
                db.session.delete(log)
            logger.info(f"Deleted {len(classification_logs)} classification logs for CV {cv_id}")
        
        # Xóa tất cả CVProcessingQueue records liên quan
        queue_items = CVProcessingQueue.query.filter_by(cv_id=cv_id).all()
        if queue_items:
            for queue_item in queue_items:
                db.session.delete(queue_item)
            logger.info(f"Deleted {len(queue_items)} queue items for CV {cv_id}")
        
        # Xóa file vật lý
        file_path = os.path.join(UPLOAD_FOLDER, cv.file_name)
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info(f"Deleted physical file: {file_path}")
            except Exception as file_error:
                logger.warning(f"Error deleting physical file {file_path}: {str(file_error)}")
                # Continue even if file deletion fails
        
        # Xóa record CV từ database
        db.session.delete(cv)
        db.session.commit()
        
        logger.info(f"Successfully deleted CV {cv_id} for user {current_user_id}")
        
        return jsonify({
            'success': True,
            'message': 'CV deleted successfully'
        }), 200
    
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting CV {cv_id}: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Error deleting CV: {str(e)}'
        }), 500


@cv_bp.route('/api/upload', methods=['POST'])
@jwt_required()
def api_upload_cv():
    """API endpoint để upload CV với JWT authentication"""
    try:
        current_user_id = get_user_id_from_jwt()
        
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'message': 'No file provided'
            }), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({
                'success': False,
                'message': 'No file selected'
            }), 400
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            os.makedirs(UPLOAD_FOLDER, exist_ok=True)
            
            # Lưu file và lấy kích thước
            file.save(file_path)
            file_size = os.path.getsize(file_path)
            file_type = filename.rsplit('.', 1)[1].lower()

            # Extract text from file
            # For images, skip OCR if file is too large or processing takes too long
            extracted_text = None
            try:
                # Check file size before OCR for images
                if file_type in ['jpg', 'jpeg', 'png']:
                    file_size_mb = file_size / (1024 * 1024)
                    if file_size_mb > 5:
                        logger.info(f"Image {filename} is large ({file_size_mb:.1f}MB). Skipping OCR to prevent hanging. File uploaded successfully.")
                        extracted_text = None
                    else:
                        logger.info(f"Processing image {filename} with OCR (may take 10-30 seconds)...")
                        try:
                            extracted_text = extract_text_from_file(file_path, file_type)
                        except Exception as ocr_error:
                            logger.warning(f"OCR failed for {filename}: {str(ocr_error)}. File uploaded but text extraction skipped.")
                            extracted_text = None
                else:
                    extracted_text = extract_text_from_file(file_path, file_type)
                
                if extracted_text:
                    logger.info(f"Successfully extracted text from {filename} ({len(extracted_text)} characters)")
                else:
                    if file_type in ['jpg', 'jpeg', 'png']:
                        logger.info(f"Image {filename} uploaded successfully. Text extraction skipped (this is normal for large images).")
                    else:
                        logger.warning(f"No text extracted from {filename}")
            except Exception as e:
                logger.error(f"Error extracting text from {filename}: {str(e)}")
                # Continue even if extraction fails - file is still saved

            # Classify CV if text was extracted
            predicted_category_id = None
            confidence = None
            if extracted_text:
                try:
                    category_name, confidence_score = classify_cv_by_keywords(extracted_text)
                    if category_name:
                        predicted_category_id = get_category_id_by_name(category_name, JobCategory)
                        if predicted_category_id:
                            logger.info(f"Classified CV as '{category_name}' (ID: {predicted_category_id}) with confidence {confidence_score:.2%}")
                        else:
                            logger.warning(f"Category '{category_name}' not found in database")
                        confidence = confidence_score
                except Exception as e:
                    logger.error(f"Error classifying CV: {str(e)}")
                    # Continue even if classification fails

            new_cv = CV(
                file_name=filename,
                file_type=file_type,
                file_size=file_size,
                user_id=current_user_id,
                file_content=extracted_text,
                predicted_category_id=predicted_category_id
            )
            db.session.add(new_cv)
            db.session.commit()
            
            # Create classification log if classification was successful
            if predicted_category_id and confidence is not None:
                try:
                    classification_log = ClassificationLog(
                        cv_id=new_cv.id,
                        predicted_category_id=predicted_category_id,
                        confidence=confidence,
                        user_id=current_user_id
                    )
                    db.session.add(classification_log)
                    db.session.commit()
                    logger.info(f"Created classification log for CV {new_cv.id}")
                except Exception as e:
                    logger.error(f"Error creating classification log: {str(e)}")
                    db.session.rollback()
            
            return jsonify({
                'success': True,
                'message': 'CV uploaded successfully',
                'cv': {
                    'id': new_cv.id,
                    'file_name': new_cv.file_name,
                    'file_type': new_cv.file_type,
                    'file_size': new_cv.file_size,
                    'uploaded_at': new_cv.uploaded_at.isoformat() if new_cv.uploaded_at else None
                }
            }), 201
        else:
            return jsonify({
                'success': False,
                'message': 'File type not allowed. Chỉ hỗ trợ: PDF, DOCX, TXT, JPG, JPEG, PNG'
            }), 400
    
    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'message': f'Error uploading CV: {str(e)}'
        }), 500
