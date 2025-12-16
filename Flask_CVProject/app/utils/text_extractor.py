"""
Text extraction utilities for CV files
Supports PDF, DOCX, TXT, and Image (JPG, PNG, JPEG) file formats
"""

import os
import logging
from typing import Optional

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    Extract text from PDF file.
    Tries pdfplumber first (better), falls back to PyPDF2.
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Extracted text as string, or None if error
    """
    if not os.path.exists(file_path):
        logger.error(f"PDF file not found: {file_path}")
        return None
    
    text = ""
    
    # Try pdfplumber first (better accuracy)
    if HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(file_path) as pdf:
                pages_text = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                text = "\n".join(pages_text)
                if text:
                    logger.info(f"Successfully extracted text from PDF using pdfplumber: {file_path}")
                    return text.strip()
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed for {file_path}: {str(e)}, trying PyPDF2")
    
    # Fallback to PyPDF2
    if HAS_PYPDF2:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pages_text = []
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        try:
                            if isinstance(page_text, bytes):
                                page_text = page_text.decode('utf-8', errors='ignore')
                            pages_text.append(page_text)
                        except Exception as e:
                            logger.warning(f"Error decoding page text: {str(e)}")
                            pages_text.append(page_text)
                text = "\n".join(pages_text)
                if text:
                    logger.info(f"Successfully extracted text from PDF using PyPDF2: {file_path}")
                    return text.strip()
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed for {file_path}: {str(e)}")
            return None
    else:
        logger.error("No PDF extraction library available (pdfplumber or PyPDF2)")
        return None
    
    return None


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Extract text from DOCX file.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text as string, or None if error
    """
    if not os.path.exists(file_path):
        logger.error(f"DOCX file not found: {file_path}")
        return None
    
    if not HAS_DOCX:
        logger.error("python-docx library not installed")
        return None
    
    try:
        doc = Document(file_path)
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        text = "\n".join(paragraphs)
        
        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        text = "\n".join(paragraphs)
        
        if text:
            logger.info(f"Successfully extracted text from DOCX: {file_path}")
            return text.strip()
        else:
            logger.warning(f"No text extracted from DOCX: {file_path}")
            return None
            
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {str(e)}")
        return None


def extract_text_from_txt(file_path: str) -> Optional[str]:
    """
    Extract text from TXT file.
    
    Args:
        file_path: Path to TXT file
        
    Returns:
        Extracted text as string, or None if error
    """
    if not os.path.exists(file_path):
        logger.error(f"TXT file not found: {file_path}")
        return None
    
    encodings = ['utf-8', 'utf-8-sig', 'cp1252', 'latin-1', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                text = file.read()
                if text:
                    logger.info(f"Successfully extracted text from TXT ({encoding}): {file_path}")
                    return text.strip()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.warning(f"Error reading TXT with {encoding}: {str(e)}")
            continue
    
    logger.error(f"TXT extraction failed for {file_path}: Could not decode with any encoding")
    return None
    
    return None


def extract_text_from_file(file_path: str, file_type: str = None) -> Optional[str]:
    """
    Extract text from file based on file type.
    Supports PDF, DOCX, TXT, and Image (JPG, PNG, JPEG) formats.
    
    Args:
        file_path: Path to file
        file_type: File extension (pdf, docx, txt, jpg, jpeg, png). If None, auto-detect from file_path
        
    Returns:
        Extracted text as string, or None if error or unsupported format
    """
    if not file_type:
        # Auto-detect from file extension
        _, ext = os.path.splitext(file_path)
        file_type = ext.lstrip('.').lower()
    
    file_type = file_type.lower()
    
    # Document formats
    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type in ['docx', 'doc']:
        return extract_text_from_docx(file_path)
    elif file_type == 'txt':
        return extract_text_from_txt(file_path)
    # Image formats (OCR)
    elif file_type in ['jpg', 'jpeg', 'png']:
        try:
            from app.utils.image_extractor import extract_text_from_image_file
            return extract_text_from_image_file(file_path, file_type)
        except ImportError:
            logger.error("Image extractor not available. Install easyocr and Pillow to enable image OCR.")
            return None
    else:
        logger.error(f"Unsupported file type: {file_type}")
        return None

